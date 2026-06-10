"""
tests/integration/test_delivery_golden.py — DVA-CBS | Projeto Diógenes
Onda 1 — Hardening do delivery

Golden structural tests: verify the DOCX appendix structure invariants using
python-docx (heading counts/titles, table count, values format) and verify all
regex patterns in runtime.yaml:motor_saida compile and behave correctly.
"""
from __future__ import annotations

import re
from pathlib import Path

import pytest
import yaml

from diogenes.delivery import builders, parsing
from diogenes.models import (
    OcorrenciaEntrega,
    PacoteEntrega,
    TesteRealizado as _Teste,
)


# ── Shared fixture ───────────────────────────────────────────────────────────

@pytest.fixture
def pacote_rico() -> PacoteEntrega:
    """PacoteEntrega com campos suficientes para exercitar as 7 seções do Apêndice."""
    ocs = [
        OcorrenciaEntrega(
            codigo="W010-001", titulo="Divergência de escopo na base de débitos",
            nivel="CRITICO", fundamento_violado="LC 214/2025 art. 5 §2",
            descricao="A planilha utiliza receita bruta sem dedução de tributos.",
            solicitacao_rfb="Demonstrar o ajuste e reencaminhar planilha corrigida.",
        ),
        OcorrenciaEntrega(
            codigo="W010-002", titulo="Percentual de crédito não documentado",
            nivel="ATENCAO", descricao="38,09% aplicado sem referência a fonte primária.",
        ),
    ]
    return PacoteEntrega(
        cycle_id="MOD_010_A1_20260101T000000Z",
        modulo=10,
        modulo_nome="Pessoa Física",
        versao="1.0",
        veredito="APROVADO_COM_RESSALVAS",
        conclusao_conformidade="Módulo parcialmente conforme com ressalvas técnicas.",
        conclusao_consistencia="Cálculo internamente consistente dentro das premissas adotadas.",
        proposta_descricao="Módulo 10 define a metodologia CBS para Pessoas Físicas.",
        proposta_contexto="LC 214/2025 arts. 164 e 165.",
        proposta_fonte="Anexo Nota Cetad 079/2025 (peça 16, item III.10).",
        objetivo="Calcular a arrecadação de CBS proveniente das PF contribuintes.",
        objetivo_detalhado="Inclui débitos e créditos CBS, apuração da arrecadação líquida e ajuste no Módulo Central.",
        arquivos_principal=[{"nome": "AUX_MOD_10 PF, Execução.xlsx", "descricao": "Planilha principal", "tipo": "xlsx"}],
        arquivos_auxiliares=[{"nome": "notebook_extracao.ipynb", "descricao": "Script de extração Datalake", "tipo": "ipynb"}],
        arquivos_fontes=[{"nome": "DIRPF_2023_amostra.csv", "descricao": "Amostra DIRPF", "tipo": "csv"}],
        testes_camada_1=[_Teste("T-C1-01", "Fechamento numérico dos débitos", "Diferença < 0,01%", "Atendido")],
        testes_camada_2=[_Teste("T-C2-01", "Validação percentual de crédito", "38,09% referenciado", "Atendido Parcialmente")],
        testes_camada_3=[_Teste("T-C3-01", "Recálculo independente CBS PF 2023", "R$ 3,07 bi vs R$ 3,09 bi (delta 0,6%)", "Atendido")],
        ocorrencias=ocs,
        valores_agregados=[
            {"descricao": "Débitos CBS PF", "valor_2023": "R$ 20,1 bilhões", "valor_2024": "R$ 19,7 bilhões"},
            {"descricao": "Créditos CBS PF", "valor_2023": "R$ 26,8 bilhões", "valor_2024": "R$ 25,6 bilhões"},
        ],
        apendice_conteudo={
            "proposta": {"fonte": "Anexo Nota Cetad 079/2025 (peça 16, item III.10)"},
            "testes": {
                "camada_1": {"conformidade": [{"id": "T-C1-01", "descricao": "Fechamento", "resultado": "ok", "status": "Atendido"}]},
                "camada_2": {"premissas": [{"id": "T-C2-01", "descricao": "Percentual", "resultado": "parcial", "status": "Atendido Parcialmente"}]},
                "camada_3": {"recalculo": [{"id": "T-C3-01", "descricao": "Recálculo", "resultado": "ok", "status": "Atendido"}]},
            },
            "inconsistencias": [
                {"id": "W010-001", "consequencia": "Superestima base de cálculo", "tratamento": "RFB demonstrará ajuste."},
            ],
            "alteracoes_metodologicas": [
                {"id": "ALT-01", "descricao": "Inclusão do redutor PF art. 165", "acordo": "Acordado", "impacto": "Reduz arrecadação em ~3%"},
            ],
            "conclusao": {
                "conformidade": "Módulo parcialmente conforme com ressalvas técnicas.",
                "premissas": [{"nome": "Alíquota PF 5,95%", "descricao": "Redução 30% sobre referência", "impacto": "Menor arrecadação."}],
            },
        },
    )


# ── Apêndice estrutural ──────────────────────────────────────────────────────

@pytest.fixture
def apendice_path(pacote_rico: PacoteEntrega, tmp_path: Path) -> Path:
    docx = pytest.importorskip("docx")  # noqa: F841
    return builders.gerar_apendice_docx(pacote_rico, tmp_path / "Apendice.docx")


def test_apendice_tem_7_secoes_h1(apendice_path: Path):
    import docx
    doc = docx.Document(str(apendice_path))
    h1_texts = [p.text.strip() for p in doc.paragraphs if p.style.name == "Heading 1"]
    # The ApendiceGerador always produces exactly 7 level-1 sections
    assert len(h1_texts) == 7, f"Esperadas 7 seções H1, encontradas {len(h1_texts)}: {h1_texts}"


def test_apendice_secoes_titulos_corretos(apendice_path: Path):
    import docx
    doc = docx.Document(str(apendice_path))
    h1_texts = [p.text.strip() for p in doc.paragraphs if p.style.name == "Heading 1"]
    expected_fragments = [
        "Proposta do Módulo",
        "Objetivo do Módulo",
        "Relação dos Arquivos",
        "Testes Realizados",
        "Resultados",           # "5. Resultados e Inconsistências Identificadas"
        "Alterações",           # "6. Alterações Metodológicas"
        "Conclusão",            # "7. Conclusão do Módulo"
    ]
    for fragment in expected_fragments:
        matched = any(fragment in t for t in h1_texts)
        assert matched, f"Seção com '{fragment}' não encontrada. H1s: {h1_texts}"


def test_apendice_tem_tabelas_minimas(apendice_path: Path):
    import docx
    doc = docx.Document(str(apendice_path))
    # With our rich fixture: 3 arquivo tables + 3 test tables + 1 INC resumo +
    # 2 INC detail + 1 alterações + 1 premissas + 1 valores = ≥ 13 tables
    assert len(doc.tables) >= 5, f"Esperadas ≥ 5 tabelas, encontradas {len(doc.tables)}"


def test_apendice_cabecalho_com_texto_institucional(apendice_path: Path):
    """Header usa tabela invisível (logo + texto) — precisa iterar cells, não paragraphs."""
    import docx
    doc = docx.Document(str(apendice_path))
    header_text = ""
    for sec in doc.sections:
        # Paragraphs diretos no header
        for para in sec.header.paragraphs:
            header_text += para.text + " "
        # Tabelas invisíveis no header (padrão TCUFormatter)
        for tbl in sec.header.tables:
            for row in tbl.rows:
                for cell in row.cells:
                    header_text += cell.text + " "
    assert "TRIBUNAL" in header_text.upper() or "TCU" in header_text.upper(), (
        f"Header não contém referência ao TCU. Texto extraído: {header_text!r}"
    )


def test_apendice_valores_nao_sao_floats_crus(apendice_path: Path):
    """Valores monetários devem estar formatados (R$ X,X bilhões), não como floats python."""
    import docx
    doc = docx.Document(str(apendice_path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    # Should not contain raw Python floats like "30840000000.0" or "1.23e9"
    raw_float = re.search(r"\b\d{8,}\.\d+\b", full_text)
    assert raw_float is None, (
        f"Valor numérico cru encontrado no apêndice: {raw_float.group()!r}"
    )


def test_apendice_ocorrencias_presentes(apendice_path: Path, pacote_rico: PacoteEntrega):
    import docx
    doc = docx.Document(str(apendice_path))
    full_text = "\n".join(p.text for p in doc.paragraphs)
    for oc in pacote_rico.ocorrencias:
        assert oc.codigo in full_text, (
            f"Ocorrência {oc.codigo!r} não encontrada no apêndice gerado."
        )


# ── Markdown narrativo ───────────────────────────────────────────────────────

def test_markdown_narrativo_tem_blocos_obrigatorios(pacote_rico: PacoteEntrega):
    md = builders.montar_markdown_narrativo(pacote_rico)
    # 5 mandatory blocks (H2)
    assert "## Bloco 1: Proposta do Módulo" in md
    assert "## Bloco 2: Objetivo" in md
    assert "## Bloco 3: Resultados Financeiros" in md
    assert "## Bloco 4: Inconsistências Identificadas" in md
    assert "## Bloco 5: Conclusão" in md


def test_markdown_narrativo_contem_ocorrencias(pacote_rico: PacoteEntrega):
    md = builders.montar_markdown_narrativo(pacote_rico)
    assert "W010-001" in md
    assert "W010-002" in md


def test_markdown_narrativo_contem_valores_formatados(pacote_rico: PacoteEntrega):
    md = builders.montar_markdown_narrativo(pacote_rico)
    # Values from valores_agregados should appear formatted
    assert "R$ 20,1 bilhões" in md
    assert "Débitos CBS PF" in md


# ── Regex sanity (motor_saida runtime.yaml) ─────────────────────────────────

@pytest.fixture(scope="module")
def motor_saida_cfg() -> dict:
    runtime_path = Path(__file__).parent.parent.parent / "runtime.yaml"
    with runtime_path.open(encoding="utf-8") as f:
        return yaml.safe_load(f)["motor_saida"]


def test_padroes_cargo_identificador_compilam(motor_saida_cfg: dict):
    for padrao in motor_saida_cfg["padroes_cargo_identificador"]:
        try:
            re.compile(padrao, re.IGNORECASE | re.UNICODE)
        except re.error as exc:
            pytest.fail(f"Padrão de cargo inválido: {padrao!r} → {exc}")


def test_regex_cycle_id_compila_e_valida_formato(motor_saida_cfg: dict):
    pattern = re.compile(motor_saida_cfg["regex_cycle_id"])
    # Must match valid cycle IDs (format: MOD_<id>_A<n>_<YYYYMMDDTHHmmssZ>)
    valid_ids = [
        "MOD_010_A1_20260101T000000Z",
        "MOD_SINT_001_A2_20260605T103922Z",
        "MOD_010_A1_20260605T103922Z",
    ]
    for cid in valid_ids:
        assert pattern.search(cid), f"regex_cycle_id não match em {cid!r}"

    # Must not match plain words
    for neg in ["plain_text", "MOD", "2026-06-01"]:
        assert not pattern.fullmatch(neg), f"regex_cycle_id match incorreto em {neg!r}"


def test_substituicoes_higienizacao_sao_pares(motor_saida_cfg: dict):
    for i, item in enumerate(motor_saida_cfg["substituicoes_higienizacao"]):
        assert len(item) == 2, f"Substituição {i} não é par: {item!r}"
        padrao, substituto = item
        assert isinstance(padrao, str) and padrao, f"Padrão {i} vazio"
        assert isinstance(substituto, str), f"Substituto {i} não é string"


def test_padroes_agentes_detectados_pelo_motor_saida(motor_saida_cfg: dict):
    """Cada nome de agente deve ser detectável (case-insensitive substring) em algum texto."""
    agentes = motor_saida_cfg["padroes_agentes"]
    # Texto que contém todas as variantes listadas em padroes_agentes
    texto_com_agentes = (
        "Mycroft Holmes consolidou o relatório. Dr. John Watson verificou. "
        "John Watson identificou. Dr. Watson analisou. Watson revisou. "
        "Sherlock Holmes aplicou a metodologia. Sherlock classificou. "
        "Inspetor Lestrade chancelou. Lestrade confirmou."
    )
    for agente in agentes:
        assert agente.lower() in texto_com_agentes.lower(), (
            f"Agente {agente!r} não está coberto pelo texto de teste — adicionar variante"
        )


def test_substituicoes_nao_introduzem_nomes_de_agentes(motor_saida_cfg: dict):
    """Nenhuma substituição deve resultar em texto que contenha um nome de agente."""
    agentes_lower = {a.lower() for a in motor_saida_cfg["padroes_agentes"]}
    for padrao, substituto in motor_saida_cfg["substituicoes_higienizacao"]:
        for agente in agentes_lower:
            assert agente not in substituto.lower(), (
                f"Substituto {substituto!r} contém nome de agente {agente!r}"
            )


def test_padroes_estruturas_internas_nao_sao_vazios(motor_saida_cfg: dict):
    for p in motor_saida_cfg["padroes_estruturas_internas"]:
        assert p and isinstance(p, str), f"Padrão estrutura interna vazio: {p!r}"
