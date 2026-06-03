"""
tests/unit/test_file_prep.py — DVA-CBS | Projeto Diógenes
Etapa 2: tetos de truncamento recalibrados + estratégia cabeça+cauda.
"""
from __future__ import annotations

import json

from diogenes.agents import file_prep
from diogenes.agents.file_prep import preparar_arquivo

# ── SQL ──────────────────────────────────────────────────────────────────────

def test_sql_medio_passa_inteiro(tmp_path):
    # ~24k chars (tamanho real do MOD_010) deve passar inteiro (era cortado em 8k)
    raw = "SELECT col_a, col_b FROM tabela_x WHERE ano = 2023;\n" * 460
    assert len(raw) > 8_000
    p = tmp_path / "q.sql"
    p.write_text(raw, encoding="utf-8")
    out = preparar_arquivo(p)
    assert "[TRUNCADO" not in out


def test_sql_gigante_trunca_no_teto(tmp_path):
    raw = "SELECT 1;\n" * 8000  # ~80k chars
    p = tmp_path / "big.sql"
    p.write_text(raw, encoding="utf-8")
    out = preparar_arquivo(p)
    assert "[TRUNCADO" in out
    assert len(out) <= file_prep._SQL_MAX_CHARS + 100


# ── CSV (cabeça + cauda + resumo) ────────────────────────────────────────────

def test_csv_pequeno_inteiro(tmp_path):
    p = tmp_path / "small.csv"
    p.write_text("id;v\n1;10\n2;20\n", encoding="utf-8")
    out = preparar_arquivo(p)
    assert out.strip().endswith("2;20")
    assert "omitidas" not in out


def test_csv_grande_cabeca_cauda_e_total(tmp_path):
    linhas = ["id;valor"] + [f"{i};{i*10}" for i in range(1, 5001)]
    linhas.append("TOTAL;999999")  # última linha = total (deve aparecer)
    p = tmp_path / "big.csv"
    p.write_text("\n".join(linhas), encoding="utf-8")
    out = preparar_arquivo(p)
    assert "TOTAL;999999" in out                 # cauda preservada
    assert "id;valor" in out                      # cabeçalho preservado
    assert "linhas omitidas" in out
    assert f"{file_prep._CSV_HEAD_LINHAS}" in out  # resumo cita o teto


# ── Notebook / texto genérico ────────────────────────────────────────────────

def test_notebook_trunca_no_teto(tmp_path):
    big_src = "x = 1\n" * 30_000  # ~180k chars de código
    nb = {
        "cells": [{"cell_type": "code", "source": big_src, "metadata": {},
                   "outputs": [], "execution_count": None}],
        "metadata": {}, "nbformat": 4, "nbformat_minor": 5,
    }
    p = tmp_path / "big.ipynb"
    p.write_text(json.dumps(nb), encoding="utf-8")
    out = preparar_arquivo(p)
    assert "[TRUNCADO" in out
    assert len(out) <= file_prep._NOTEBOOK_MAX_CHARS + 100


def test_texto_generico_trunca(tmp_path):
    p = tmp_path / "grande.py"
    p.write_text("a = 0\n" * 20_000, encoding="utf-8")
    out = preparar_arquivo(p)
    assert "[TRUNCADO" in out
    assert len(out) <= file_prep._TEXTO_MAX_CHARS + 100


# ── XLSX (cabeça + cauda por aba + fórmulas + teto global) ───────────────────

def test_xlsx_cabeca_cauda_e_formula(tmp_path):
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Apuracao"
    ws.append(["id", "valor"])             # cabeçalho
    for i in range(1, 300):
        ws.append([i, i * 10])
    ws["C1"] = "=SUM(B2:B300)"             # fórmula deve ser preservada
    ws.append(["TOTAL_FINAL", 99999])      # última linha (cauda)
    p = tmp_path / "plan.xlsx"
    wb.save(p)
    out = preparar_arquivo(p)
    assert "[FORMULA:=SUM(B2:B300)]" in out
    assert "TOTAL_FINAL" in out            # cauda preservada
    assert "linhas omitidas" in out        # corpo do meio omitido
    assert "Aba: Apuracao" in out


# ── DOCX / DOC ───────────────────────────────────────────────────────────────

def test_docx_leitura_conteudo_e_tabela(tmp_path):
    import docx
    doc = docx.Document()
    doc.add_paragraph("Este é um parágrafo de teste para o Diógenes.")
    
    table = doc.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "Cabeçalho 1"
    table.cell(0, 1).text = "Cabeçalho 2"
    table.cell(1, 0).text = "Dado 1"
    table.cell(1, 1).text = "Dado 2"
    
    p = tmp_path / "doc_teste.docx"
    doc.save(p)
    
    out = preparar_arquivo(p)
    assert "Este é um parágrafo de teste para o Diógenes." in out
    assert "Cabeçalho 1\tCabeçalho 2" in out
    assert "Dado 1\tDado 2" in out


def test_docx_fallback_docling_se_invalido(tmp_path, monkeypatch):
    # Simula um arquivo .doc legado (que fará python-docx falhar e cair no docling)
    p = tmp_path / "legado.doc"
    p.write_text("Conteúdo binário simulado", encoding="utf-8")
    
    # Mock docling converter para não precisar rodar conversão de modelo neural em teste unitário
    class MockDocument:
        def export_to_markdown(self):
            return "Conteúdo convertido via Docling"
            
    class MockResult:
        def __init__(self):
            self.document = MockDocument()

    class MockDocumentConverter:
        def convert(self, path_str):
            return MockResult()

    monkeypatch.setattr("docling.document_converter.DocumentConverter", MockDocumentConverter)
    
    out = preparar_arquivo(p)
    assert "Conteúdo convertido via Docling" in out

