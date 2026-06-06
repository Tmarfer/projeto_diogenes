"""
agents/file_prep.py — Preparação de conteúdo de arquivos para Watson.
Extrai conteúdo legível de xlsx, sql, ipynb, pdf, md.
Referência normativa: RF-WA-03 a RF-WA-05 (PRD v0.1), Bloco 9.3.2 (SDD v0.1)
"""
from __future__ import annotations

from pathlib import Path

# ---------------------------------------------------------------------------
# Tetos de truncamento (Etapa 2 — recalibrados para o pack real do MOD_010).
# file_prep é a ÚNICA fonte de truncamento no fluxo oficial (Watson usa
# preparar_arquivo direto). Estratégia para tabular grande: cabeça + cauda +
# resumo, preservando cabeçalho/amostra E os totais que ficam no fim.
# ---------------------------------------------------------------------------
_SQL_MAX_CHARS = 40_000        # SQLs reais (~24k) passam inteiros
_CSV_HEAD_LINHAS = 1_500       # CSV: primeiras N linhas (cabeçalho + amostra)
_CSV_TAIL_LINHAS = 150         # CSV: últimas M linhas (totais/fechamentos)
_CSV_MAX_CHARS = 120_000       # CSV: teto de chars (CSVs largos têm linhas longas)
_XLSX_LINHAS_ABA = 80          # XLSX: cabeça por aba (mantém [FORMULA:...])
_XLSX_TAIL_LINHAS = 20         # XLSX: cauda por aba (totais)
_XLSX_MAX_CHARS = 60_000       # XLSX: teto global por arquivo (multi-aba)
_NOTEBOOK_MAX_CHARS = 80_000   # Notebook: código extraído
_TEXTO_MAX_CHARS = 60_000      # genérico p/ .txt/.md/.py e demais textos

# ---------------------------------------------------------------------------
# Classificação de arquivos da entrega (fonte única — reusada por bench e
# pelo Motor de Start). Separa o que é analisável (cat. A+B+docs) do que é
# metadado/manifesto de entrega (cat. E) e logs.
# ---------------------------------------------------------------------------

# Extensões aceitas para análise (convertidas via preparar_arquivo).
EXTENSOES_ANALISE: set[str] = {
    ".xlsx", ".csv", ".sql", ".ipynb", ".py", ".pdf", ".md", ".txt", ".docx", ".doc",
}
# Diretórios ignorados na varredura da entrega (logs, metadados, catálogos).
DIRS_IGNORADOS: set[str] = {
    "03_LOGS", "_LOGS", "02_INSUMOS_GERADOS", "_CATALOGO",
    "progress", "_runtime", "__pycache__",
}
# Arquivos de metadados / manifesto de entrega — nunca analisados como conteúdo.
ARQUIVOS_IGNORADOS: set[str] = {
    "catalogo.json", "metadados.json", "hashes_integridade.txt",
    "intake.log",
}
# Arquivos de inventário/protocolo que DEVEM ser analisados mesmo estando sob um
# diretório de metadados (ex.: 02_INSUMOS_GERADOS). Os agentes precisam saber o que
# foi de fato entregue, para organizar a redação e citar a origem dos achados.
# A allowlist tem precedência sobre DIRS_IGNORADOS e ARQUIVOS_IGNORADOS.
ARQUIVOS_SEMPRE_ANALISE: set[str] = {
    "protocolo_recebimento.md", "inventario.xlsx",
}

# Extensão → rótulo de tipo legível.
_TIPO_POR_EXTENSAO: dict[str, str] = {
    ".csv": "csv",
    ".xlsx": "planilha",
    ".xls": "planilha",
    ".sql": "sql",
    ".ipynb": "notebook",
    ".py": "script",
    ".pdf": "documento",
    ".md": "documento",
    ".txt": "documento",
    ".docx": "documento",
    ".doc": "documento",
    ".json": "esquema",
}


def rotulo_tipo(filename: str) -> str:
    """Rótulo de tipo do arquivo a partir da extensão (csv, sql, notebook, ...)."""
    return _TIPO_POR_EXTENSAO.get(Path(filename).suffix.lower(), "documento")


def classificar(rel_path: Path | str) -> tuple[bool, str]:
    """Classifica um arquivo da entrega.

    Retorna (analisavel, tipo): `analisavel` é False para arquivos sob diretórios
    da denylist, arquivos de metadados, ou extensões fora da allowlist. `tipo` é o
    rótulo legível (ver rotulo_tipo); para não-analisáveis, "metadados".
    """
    p = Path(rel_path)
    # Allowlist tem precedência: inventário/protocolo são analisados mesmo sob
    # diretórios de metadados, desde que tenham extensão analisável.
    if p.name.lower() in ARQUIVOS_SEMPRE_ANALISE and p.suffix.lower() in EXTENSOES_ANALISE:
        return True, rotulo_tipo(p.name)
    if any(parte in DIRS_IGNORADOS for parte in p.parts[:-1]):
        return False, "metadados"
    if p.name.lower() in ARQUIVOS_IGNORADOS:
        return False, "metadados"
    if p.suffix.lower() not in EXTENSOES_ANALISE:
        return False, "metadados"
    return True, rotulo_tipo(p.name)


def preparar_arquivo(path: Path, perfil_dir: Path | None = None) -> str:
    """Retorna representação textual do arquivo para inclusão em prompt.

    Se perfil_dir for fornecido, tenta prefixar com o perfil analítico
    gerado pelo MotorPerfilamento (.perfil.yaml). O prefixo fornece a
    Watson estatísticas completas (unicidade, nulos, ranges) independente
    do truncamento da amostra textual.
    """
    prefixo = ""
    if perfil_dir is not None:
        try:
            from diogenes.motors.motor_perfilamento import ler_perfil_texto
            texto_perfil = ler_perfil_texto(Path(perfil_dir), path)
            if texto_perfil:
                prefixo = f"[PERFIL ANALÍTICO — {path.name}]\n{texto_perfil}\n\n"
        except Exception:  # noqa: BLE001
            pass  # falha silenciosa: arquivo sem perfil não bloqueia Watson

    ext = path.suffix.lower()
    if ext == ".xlsx":
        conteudo = _ler_excel(path)
    elif ext == ".csv":
        conteudo = _ler_csv(path)
    elif ext == ".sql":
        conteudo = _ler_sql(path)
    elif ext == ".ipynb":
        conteudo = _ler_notebook(path)
    elif ext == ".pdf":
        conteudo = _ler_pdf(path)
    elif ext in (".docx", ".doc"):
        conteudo = _ler_docx(path)
    else:
        # Markdown, scripts e texto puro — teto genérico de proteção
        try:
            conteudo = _truncar_texto(path.read_text(encoding="utf-8"), _TEXTO_MAX_CHARS)
        except Exception as e:
            conteudo = f"[ERRO AO LER ARQUIVO: {e}]"

    return prefixo + conteudo


def _truncar_texto(texto: str, teto: int) -> str:
    """Trunca texto no teto de chars com nota, preservando o início."""
    if len(texto) <= teto:
        return texto
    return texto[:teto] + f"\n\n[TRUNCADO — {len(texto)} chars totais, exibidos {teto}]"


def _ler_csv(path: Path) -> str:
    """CSV: cabeça + cauda + resumo (preserva cabeçalho/amostra e totais finais)."""
    try:
        linhas = path.read_text(encoding="utf-8", errors="replace").splitlines()
    except Exception as e:
        return f"[ERRO AO LER CSV: {e}]"
    total = len(linhas)
    if total <= _CSV_HEAD_LINHAS + _CSV_TAIL_LINHAS:
        return _truncar_texto("\n".join(linhas), _CSV_MAX_CHARS)
    cabeca = "\n".join(linhas[:_CSV_HEAD_LINHAS])
    cauda = "\n".join(linhas[-_CSV_TAIL_LINHAS:])
    omitidas = total - _CSV_HEAD_LINHAS - _CSV_TAIL_LINHAS
    # Teto de chars: a cauda (totais) é prioritária; corta a cabeça se necessário
    # (CSVs largos têm linhas longas — limitar por linhas não basta).
    orcamento_cabeca = _CSV_MAX_CHARS - len(cauda) - 300
    if len(cabeca) > orcamento_cabeca > 0:
        cabeca = cabeca[:orcamento_cabeca] + "\n[... cabeça truncada por teto de chars]"
    return (
        f"[CSV: {total} linhas — exibindo {_CSV_HEAD_LINHAS} iniciais + "
        f"{_CSV_TAIL_LINHAS} finais]\n"
        f"{cabeca}\n\n"
        f"⚠ {omitidas} linhas omitidas (meio do arquivo)\n\n"
        f"--- últimas {_CSV_TAIL_LINHAS} linhas ---\n"
        f"{cauda}"
    )


def _ler_excel(path: Path) -> str:
    try:
        import openpyxl
        wb = openpyxl.load_workbook(path, data_only=False)
        partes: list[str] = []
        total_chars = 0
        abas_omitidas = 0
        for sheet_name in wb.sheetnames:
            if total_chars >= _XLSX_MAX_CHARS:
                abas_omitidas += 1
                continue
            ws = wb[sheet_name]
            bloco = _aba_para_texto(ws, sheet_name)
            partes.append(bloco)
            total_chars += len(bloco)
        if abas_omitidas:
            partes.append(
                f"\n[⚠ {abas_omitidas} aba(s) omitida(s) — teto de "
                f"{_XLSX_MAX_CHARS} chars por arquivo atingido]"
            )
        return "\n\n".join(partes)
    except Exception as e:
        return f"[ERRO AO LER EXCEL: {e}]"


def _fmt_celula(val: object) -> str:
    if val is None:
        return ""
    if isinstance(val, str) and val.startswith("="):
        return f"[FORMULA:{val}]"
    return str(val)


def _aba_para_texto(ws, sheet_name: str) -> str:
    """Representa uma aba como cabeça + cauda de linhas (preserva fórmulas e totais)."""
    linhas = [
        "\t".join(_fmt_celula(c.value) for c in row)
        for row in ws.iter_rows(values_only=False)
    ]
    n = len(linhas)
    if n <= _XLSX_LINHAS_ABA + _XLSX_TAIL_LINHAS:
        corpo = "\n".join(linhas)
    else:
        cabeca = "\n".join(linhas[:_XLSX_LINHAS_ABA])
        cauda = "\n".join(linhas[-_XLSX_TAIL_LINHAS:])
        omitidas = n - _XLSX_LINHAS_ABA - _XLSX_TAIL_LINHAS
        corpo = (
            f"{cabeca}\n⚠ {omitidas} linhas omitidas (meio)\n"
            f"--- últimas {_XLSX_TAIL_LINHAS} linhas ---\n{cauda}"
        )
    return f"--- Aba: {sheet_name} ({n} linhas) ---\n{corpo}"


def _ler_sql(path: Path) -> str:
    try:
        import sqlparse
        raw = path.read_text(encoding="utf-8")
        formatado = sqlparse.format(raw, reindent=True, keyword_case="upper")
        if len(formatado) > _SQL_MAX_CHARS:
            return formatado[:_SQL_MAX_CHARS] + f"\n\n[TRUNCADO — {len(formatado)} chars totais]"
        return formatado
    except Exception as e:
        return f"[ERRO AO LER SQL: {e}]"


def _ler_notebook(path: Path) -> str:
    try:
        import nbformat
        nb = nbformat.read(str(path), as_version=4)
        celulas: list[str] = []
        for i, cell in enumerate(nb.cells):
            if cell.cell_type == "code":
                src = cell.source.strip()
                if src:
                    celulas.append(f"# Célula {i+1}\n{src}")
        if not celulas:
            return "[Notebook sem células de código]"
        return _truncar_texto("\n\n".join(celulas), _NOTEBOOK_MAX_CHARS)
    except Exception as e:
        return f"[ERRO AO LER NOTEBOOK: {e}]"


def _ler_pdf(path: Path) -> str:
    try:
        from pdfminer.high_level import extract_text
        texto = extract_text(str(path))
        return texto.strip() if texto else "[PDF sem texto extraível]"
    except Exception as e:
        return f"[ERRO AO LER PDF: {e}]"


def _ler_docx(path: Path) -> str:
    """Lê documentos de Word (.docx e .doc) com fallback robusto."""
    try:
        import docx
        doc = docx.Document(path)
        paragrafos = [p.text for p in doc.paragraphs]
        tabelas_texto = []
        for table in doc.tables:
            for row in table.rows:
                linha = "\t".join(cell.text.strip() for cell in row.cells)
                tabelas_texto.append(linha)
        corpo = "\n".join(paragrafos)
        if tabelas_texto:
            corpo += "\n\n--- Tabelas no Documento ---\n" + "\n".join(tabelas_texto)
        return _truncar_texto(corpo, _TEXTO_MAX_CHARS)
    except Exception as e:
        # Se falhar (ex: por ser .doc legado binário), tenta com docling
        try:
            from docling.document_converter import DocumentConverter
            converter = DocumentConverter()
            result = converter.convert(str(path))
            texto = result.document.export_to_markdown()
            return _truncar_texto(texto, _TEXTO_MAX_CHARS)
        except Exception as e2:
            return f"[ERRO AO LER DOCUMENTO: {e} | Fallback docling falhou: {e2}]"
